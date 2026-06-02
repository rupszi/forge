"""Local artifact storage + export (M7) — documents stay on disk, never cloud.

Generated documents are written under ``<project>/.forge/artifacts/``. Markdown
is the source format; ``html`` and ``txt`` export with stdlib only (no new
runtime deps). ``docx`` is best-effort via the optional ``python-docx`` package
and degrades to Markdown with a clear note when it isn't installed.
"""

from __future__ import annotations

import html as _html
import re
from pathlib import Path

_SUPPORTED = {"md", "txt", "html", "docx"}
_slug_re = re.compile(r"[^a-z0-9._-]+")


def _slugify(name: str) -> str:
    base = Path(name).name.lower().replace(" ", "-")
    slug = _slug_re.sub("-", base).strip("-.") or "document"
    return slug[:80]


def markdown_to_html(md: str) -> str:
    """Minimal, safe Markdown→HTML (stdlib only).

    Supports ATX headings, unordered lists, and paragraphs. All text is
    HTML-escaped first, so embedded markup can't inject into the output —
    important because document bodies come from an LLM.
    """
    lines = md.splitlines()
    out: list[str] = []
    in_list = False

    def close_list() -> None:
        nonlocal in_list
        if in_list:
            out.append("</ul>")
            in_list = False

    para: list[str] = []

    def flush_para() -> None:
        if para:
            out.append("<p>" + " ".join(para) + "</p>")
            para.clear()

    for raw in lines:
        line = raw.rstrip()
        esc = _html.escape(line.strip())
        heading = re.match(r"^(#{1,6})\s+(.*)$", line.strip())
        if heading:
            flush_para()
            close_list()
            level = len(heading.group(1))
            out.append(f"<h{level}>{_html.escape(heading.group(2))}</h{level}>")
        elif re.match(r"^[-*]\s+", line.strip()):
            flush_para()
            if not in_list:
                out.append("<ul>")
                in_list = True
            item = _html.escape(re.sub(r"^[-*]\s+", "", line.strip()))
            out.append(f"<li>{item}</li>")
        elif not line.strip():
            flush_para()
            close_list()
        else:
            para.append(esc)
    flush_para()
    close_list()
    body = "\n".join(out)
    return f"<!doctype html>\n<html>\n<head><meta charset='utf-8'></head>\n<body>\n{body}\n</body>\n</html>\n"


def artifacts_dir(base_path: str = ".") -> Path:
    d = Path(base_path) / ".forge" / "artifacts"
    d.mkdir(parents=True, exist_ok=True)
    return d


def save_artifact(name: str, content: str, fmt: str = "md", base_path: str = ".") -> str:
    """Write ``content`` to ``.forge/artifacts/<slug>.<fmt>``. Returns the path."""
    if fmt not in _SUPPORTED:
        msg = f"unsupported artifact format {fmt!r} (supported: {sorted(_SUPPORTED)})"
        raise ValueError(msg)

    slug = _slugify(name)
    out_dir = artifacts_dir(base_path)

    if fmt == "html":
        data = markdown_to_html(content)
    elif fmt == "docx":
        return _save_docx(out_dir / f"{slug}.docx", content)
    else:  # md / txt
        data = content

    path = out_dir / f"{slug}.{fmt}"
    path.write_text(data, encoding="utf-8")
    return str(path)


def _save_docx(path: Path, markdown: str) -> str:
    """Best-effort .docx via python-docx; falls back to .md if unavailable."""
    try:
        from docx import Document  # type: ignore[import-not-found]
    except ImportError:
        fallback = path.with_suffix(".md")
        fallback.write_text(
            markdown + "\n\n<!-- docx export needs `pip install python-docx` -->\n",
            encoding="utf-8",
        )
        return str(fallback)

    doc = Document()
    for line in markdown.splitlines():
        s = line.strip()
        heading = re.match(r"^(#{1,6})\s+(.*)$", s)
        if heading:
            doc.add_heading(heading.group(2), level=min(len(heading.group(1)), 9))
        elif re.match(r"^[-*]\s+", s):
            doc.add_paragraph(re.sub(r"^[-*]\s+", "", s), style="List Bullet")
        elif s:
            doc.add_paragraph(s)
    doc.save(str(path))
    return str(path)
